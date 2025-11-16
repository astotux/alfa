import fitz
from docx import Document as DocxDocument
from services.llm_service import get_llm_response
import logging
import os
import re
import unicodedata
import asyncio
import time
from typing import Optional, Callable

logger = logging.getLogger(__name__)

# Глобальные ограничения
MAX_PDF_PAGES = 20
MAX_DOCX_PARAGRAPHS = 500
MAX_TEXT_LENGTH = 15000
MAX_TOTAL_TIME = 310  # 5 минут + запас
LLM_TIMEOUT = 300  # 5 минут на запрос к ИИ


def clean_text(text: str) -> str:
    """Базовая очистка текста с сохранением читаемости"""
    if not text:
        return ""
    text = unicodedata.normalize('NFKC', text)
    text = ''.join(c for c in text if c.isprintable() and ord(c) < 65536)
    text = re.sub(r'[^\S\r\n]+', ' ', text)
    text = re.sub(r'[\r\n]+', '\n', text)
    return text.strip()


async def extract_text_from_pdf(file_path: str, progress_callback: Optional[Callable] = None) -> str:
    """Извлечение текста из PDF с прогрессом"""
    if progress_callback:
        await progress_callback("Подготовка документа к анализу", 10)

    if not os.path.exists(file_path):
        raise FileNotFoundError("Файл не найден")

    file_size = os.path.getsize(file_path)
    if file_size > 15 * 1024 * 1024:
        raise ValueError("Размер файла превышает 15MB")
    if file_size == 0:
        raise ValueError("PDF файл пустой")

    try:
        doc = await asyncio.to_thread(fitz.open, file_path)

        if doc.is_encrypted:
            raise ValueError("PDF защищен паролем. Снимите защиту перед анализом.")
        if doc.page_count == 0:
            raise ValueError("PDF не содержит страниц")

        total_pages = min(doc.page_count, MAX_PDF_PAGES)
        full_text_parts = []

        if progress_callback:
            await progress_callback(f"Анализ страниц: 0/{total_pages}", 20)

        for page_num in range(total_pages):
            page = doc[page_num]
            try:
                text = await asyncio.to_thread(page.get_text, "text")
                text = clean_text(text)
            except Exception as e:
                logger.warning(f"Ошибка извлечения текста со страницы {page_num + 1}: {str(e)}")
                text = ""

            if len(text) > 20:
                full_text_parts.append(text)

            if progress_callback:
                progress = 20 + int((page_num + 1) / total_pages * 20)
                await progress_callback(f"Анализ страниц: {page_num + 1}/{total_pages}", progress)

        full_text = "\n\n".join(full_text_parts)[:MAX_TEXT_LENGTH]

        if not full_text or len(re.findall(r'[а-яА-ЯёЁ]', full_text)) < 10:
            raise ValueError("Не удалось извлечь текст. Возможно, это скан или защищенный PDF")

        return full_text

    finally:
        if 'doc' in locals():
            try:
                doc.close()
            except:
                pass


async def process_pdf_document(input_path: str, instruction: str, progress_callback: Optional[Callable] = None) -> str:
    """Анализ PDF документа с корректным прогресс-баром"""
    total_start = time.time()

    try:
        # Этап 1: Извлечение текста (5-40%)
        if progress_callback:
            await progress_callback("Начало обработки документа", 5)

        original_text = await extract_text_from_pdf(input_path, progress_callback)

        # Этап 2: Анализ через ИИ (40-90%)
        if progress_callback:
            await progress_callback("Передача документа на анализ ИИ", 40)

        user_request = instruction.strip() if instruction.strip() else "Провести комплексную проверку оформления и содержания документа"

        prompt = f"""
ТЕКСТ ДОКУМЕНТА:
{original_text[:3000]}

ЗАПРОС НА АНАЛИЗ: {user_request}

ПРОВЕДИТЕ ОФИЦИАЛЬНЫЙ АНАЛИЗ ДОКУМЕНТА ПО СЛЕДУЮЩИМ КРИТЕРИЯМ:

1. ОФОРМЛЕНИЕ ДОКУМЕНТА:
   - Наличие и корректность обязательных реквизитов (дата, номер, наименование организации)
   - Структура документа и логичность разделов
   - Форматирование текста (отступы, шрифты, интервалы)

2. СОДЕРЖАНИЕ ДОКУМЕНТА:
   - Грамматическая и пунктуационная правильность
   - Юридическая точность формулировок (если применимо)
   - Отсутствие противоречий и дублирования информации
   - Полнота и ясность изложения

3. РЕКОМЕНДАЦИИ:
   - Перечень конкретных замечаний с указанием разделов
   - Предложения по устранению недостатков
   - Оценка готовности документа к официальному использованию

ТРЕБОВАНИЯ К ОТВЕТУ:
- Предоставьте анализ в форме официального документа
- Используйте деловой стиль изложения
- Избегайте markdown-разметки и специальных символов форматирования
- Структурируйте ответ по разделам без использования заголовков в формате Markdown
- Укажите конкретные примеры из текста при выявлении ошибок
"""

        try:
            if progress_callback:
                await progress_callback("ИИ анализирует документ (до 5 минут)", 50)

            analysis = await asyncio.wait_for(
                get_llm_response(
                    query=prompt,
                    context="",
                    system_prompt=(
                        "Вы — официальный эксперт по документообороту. "
                        "Проводите анализ строго в соответствии с государственными стандартами. "
                        "Ответ оформляйте в виде служебной записки без использования markdown-разметки. "
                        "Используйте формальный стиль и конкретные формулировки."
                    )
                ),
                timeout=LLM_TIMEOUT
            )
        except asyncio.TimeoutError:
            elapsed = time.time() - total_start
            logger.error(f"Таймаут анализа: {elapsed:.1f} секунд")
            raise ValueError(
                "Превышено время ожидания ответа от системы анализа. "
                "Процесс анализа остановлен после 5 минут ожидания. "
                "Рекомендуется разделить документ на части или упростить запрос."
            )

        # Этап 3: Очистка результата (90-100%)
        if progress_callback:
            await progress_callback("Формирование финального отчета", 90)

        # Удаляем markdown-разметку
        analysis = re.sub(r'[*_`~#|]', '', analysis)
        analysis = re.sub(r'\[\d+\]', '', analysis)
        analysis = re.sub(r'\n{3,}', '\n\n', analysis)
        analysis = re.sub(r'^\s*[-•]\s*', '', analysis, flags=re.MULTILINE)

        if not analysis or len(analysis) < 50:
            raise ValueError("Получен некорректный результат анализа")

        # Финальный прогресс
        if progress_callback:
            await progress_callback("Анализ завершен", 100)

        total_time = time.time() - total_start
        logger.info(f"Анализ завершен за {total_time:.1f} секунд")

        return analysis.strip()

    except Exception as e:
        error_msg = str(e)
        if "паролем" in error_msg.lower() or "encrypted" in error_msg.lower():
            error_msg = "PDF документ защищен паролем. Анализ невозможен. Снимите защиту перед отправкой."
        elif "скан" in error_msg.lower() or "изображени" in error_msg.lower() or "image" in error_msg.lower():
            error_msg = "Документ содержит сканированные изображения. Требуется текстовая версия PDF для анализа."
        elif "размер" in error_msg.lower() or "больш" in error_msg.lower():
            error_msg = "Размер файла превышает допустимый лимит 15MB. Сократите объем документа."
        elif "таймаут" in error_msg.lower() or "timeout" in error_msg.lower():
            error_msg = "Превышено время анализа (5 минут). Разделите документ на части или повторите попытку позже."

        logger.error(f"Ошибка обработки: {error_msg}")
        raise ValueError(error_msg)


async def extract_text_from_docx(file_path: str, progress_callback: Optional[Callable] = None) -> str:
    """Извлечение текста из DOCX с прогрессом"""
    if progress_callback:
        await progress_callback("Подготовка документа к анализу", 10)

    if not os.path.exists(file_path):
        raise FileNotFoundError("Файл не найден")

    file_size = os.path.getsize(file_path)
    if file_size > 15 * 1024 * 1024:
        raise ValueError("Размер файла превышает 15MB")
    if file_size == 0:
        raise ValueError("DOCX файл пустой")

    try:
        doc = await asyncio.to_thread(DocxDocument, file_path)

        if progress_callback:
            await progress_callback("Извлечение текста из документа", 20)

        full_text_parts = []
        total_paragraphs = min(len(doc.paragraphs), MAX_DOCX_PARAGRAPHS)

        for para_idx, paragraph in enumerate(doc.paragraphs[:MAX_DOCX_PARAGRAPHS]):
            text = paragraph.text.strip()
            if text:
                cleaned_text = clean_text(text)
                if len(cleaned_text) > 0:
                    full_text_parts.append(cleaned_text)

            if progress_callback and para_idx % 50 == 0:
                progress = 20 + int((para_idx + 1) / total_paragraphs * 20)
                await progress_callback(f"Обработка параграфов: {para_idx + 1}/{total_paragraphs}", progress)

        # Также извлекаем текст из таблиц
        if doc.tables:
            if progress_callback:
                await progress_callback("Извлечение текста из таблиц", 35)
            
            for table_idx, table in enumerate(doc.tables):
                table_text_parts = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = clean_text(cell.text.strip())
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text_parts.append(" | ".join(row_text))
                
                if table_text_parts:
                    full_text_parts.append("\n".join(table_text_parts))

        full_text = "\n\n".join(full_text_parts)[:MAX_TEXT_LENGTH]

        if not full_text or len(re.findall(r'[а-яА-ЯёЁ]', full_text)) < 10:
            raise ValueError("Не удалось извлечь текст из документа. Возможно, документ пустой или поврежден")

        return full_text

    except Exception as e:
        error_msg = str(e).lower()
        if "corrupt" in error_msg or "damaged" in error_msg or "invalid" in error_msg:
            raise ValueError("DOCX файл поврежден или имеет неверный формат")
        raise


async def process_docx_document(input_path: str, instruction: str, progress_callback: Optional[Callable] = None) -> str:
    """Анализ DOCX документа с корректным прогресс-баром"""
    total_start = time.time()

    try:
        # Этап 1: Извлечение текста (5-40%)
        if progress_callback:
            await progress_callback("Начало обработки документа", 5)

        original_text = await extract_text_from_docx(input_path, progress_callback)

        # Этап 2: Анализ через ИИ (40-90%)
        if progress_callback:
            await progress_callback("Передача документа на анализ ИИ", 40)

        user_request = instruction.strip() if instruction.strip() else "Провести комплексную проверку оформления и содержания документа"

        prompt = f"""
ТЕКСТ ДОКУМЕНТА:
{original_text[:3000]}

ЗАПРОС НА АНАЛИЗ: {user_request}

ПРОВЕДИТЕ ОФИЦИАЛЬНЫЙ АНАЛИЗ ДОКУМЕНТА ПО СЛЕДУЮЩИМ КРИТЕРИЯМ:

1. ОФОРМЛЕНИЕ ДОКУМЕНТА:
   - Наличие и корректность обязательных реквизитов (дата, номер, наименование организации)
   - Структура документа и логичность разделов
   - Форматирование текста (отступы, шрифты, интервалы)

2. СОДЕРЖАНИЕ ДОКУМЕНТА:
   - Грамматическая и пунктуационная правильность
   - Юридическая точность формулировок (если применимо)
   - Отсутствие противоречий и дублирования информации
   - Полнота и ясность изложения

3. РЕКОМЕНДАЦИИ:
   - Перечень конкретных замечаний с указанием разделов
   - Предложения по устранению недостатков
   - Оценка готовности документа к официальному использованию

ТРЕБОВАНИЯ К ОТВЕТУ:
- Предоставьте анализ в форме официального документа
- Используйте деловой стиль изложения
- Избегайте markdown-разметки и специальных символов форматирования
- Структурируйте ответ по разделам без использования заголовков в формате Markdown
- Укажите конкретные примеры из текста при выявлении ошибок
"""

        try:
            if progress_callback:
                await progress_callback("ИИ анализирует документ (до 5 минут)", 50)

            analysis = await asyncio.wait_for(
                get_llm_response(
                    query=prompt,
                    context="",
                    system_prompt=(
                        "Вы — официальный эксперт по документообороту. "
                        "Проводите анализ строго в соответствии с государственными стандартами. "
                        "Ответ оформляйте в виде служебной записки без использования markdown-разметки. "
                        "Используйте формальный стиль и конкретные формулировки."
                    )
                ),
                timeout=LLM_TIMEOUT
            )
        except asyncio.TimeoutError:
            elapsed = time.time() - total_start
            logger.error(f"Таймаут анализа: {elapsed:.1f} секунд")
            raise ValueError(
                "Превышено время ожидания ответа от системы анализа. "
                "Процесс анализа остановлен после 5 минут ожидания. "
                "Рекомендуется разделить документ на части или упростить запрос."
            )

        # Этап 3: Очистка результата (90-100%)
        if progress_callback:
            await progress_callback("Формирование финального отчета", 90)

        # Удаляем markdown-разметку
        analysis = re.sub(r'[*_`~#|]', '', analysis)
        analysis = re.sub(r'\[\d+\]', '', analysis)
        analysis = re.sub(r'\n{3,}', '\n\n', analysis)
        analysis = re.sub(r'^\s*[-•]\s*', '', analysis, flags=re.MULTILINE)

        if not analysis or len(analysis) < 50:
            raise ValueError("Получен некорректный результат анализа")

        # Финальный прогресс
        if progress_callback:
            await progress_callback("Анализ завершен", 100)

        total_time = time.time() - total_start
        logger.info(f"Анализ завершен за {total_time:.1f} секунд")

        return analysis.strip()

    except Exception as e:
        error_msg = str(e)
        if "защищен" in error_msg.lower() or "protected" in error_msg.lower():
            error_msg = "DOCX документ защищен паролем. Анализ невозможен. Снимите защиту перед отправкой."
        elif "поврежден" in error_msg.lower() or "corrupt" in error_msg.lower() or "damaged" in error_msg.lower():
            error_msg = "DOCX файл поврежден или имеет неверный формат. Проверьте целостность файла."
        elif "размер" in error_msg.lower() or "больш" in error_msg.lower():
            error_msg = "Размер файла превышает допустимый лимит 15MB. Сократите объем документа."
        elif "таймаут" in error_msg.lower() or "timeout" in error_msg.lower():
            error_msg = "Превышено время анализа (5 минут). Разделите документ на части или повторите попытку позже."

        logger.error(f"Ошибка обработки: {error_msg}")
        raise ValueError(error_msg)